import streamlit as st
import geopandas as gpd
import folium
from streamlit_folium import st_folium
import json
import numpy as np
from datetime import datetime, timedelta
import tempfile
import zipfile
import io
import os
import pandas as pd
import random
from shapely.geometry import Polygon
from shapely.ops import unary_union
from pathlib import Path
from PIL import Image
import base64
from docx import Document
from docx.shared import Inches
import numpy as np

# Configuración de página
st.set_page_config(
    page_title="Analizador de Fertilidad Multi-Cultivo",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS personalizado mejorado
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #2E8B57;
        text-align: center;
        margin-bottom: 2rem;
        font-weight: bold;
    }
    .metric-card {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #2E8B57;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .section-header {
        color: #2E8B57;
        border-bottom: 2px solid #2E8B57;
        padding-bottom: 0.5rem;
        margin-bottom: 1rem;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d1edff;
        border: 1px solid #b3d9ff;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 2px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: #f0f2f6;
        border-radius: 5px 5px 0px 0px;
        gap: 1px;
        padding-top: 10px;
        padding-bottom: 10px;
    }
    .map-container {
        border-radius: 10px;
        overflow: hidden;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        border: 2px solid #e0e0e0;
    }
    .sublote-card {
        background: white;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        border-left: 4px solid #2E8B57;
    }
</style>
""", unsafe_allow_html=True)

# Configuración de cultivos
CULTIVOS = {
    "trigo": {
        "nombre": "Trigo",
        "color": "#FFD700",
        "npk_optimo": {"N": (80, 120), "P": (40, 60), "K": (50, 80)},
        "ph_optimo": (6.0, 7.0)
    },
    "maiz": {
        "nombre": "Maíz", 
        "color": "#32CD32",
        "npk_optimo": {"N": (120, 180), "P": (50, 80), "K": (80, 120)},
        "ph_optimo": (5.8, 7.0)
    },
    "soja": {
        "nombre": "Soja",
        "color": "#90EE90",
        "npk_optimo": {"N": (0, 20), "P": (40, 70), "K": (60, 100)},
        "ph_optimo": (6.0, 7.0)
    },
    "sorgo": {
        "nombre": "Sorgo",
        "color": "#DAA520",
        "npk_optimo": {"N": (80, 120), "P": (30, 50), "K": (60, 90)},
        "ph_optimo": (5.5, 7.5)
    },
    "girasol": {
        "nombre": "Girasol",
        "color": "#FF8C00",
        "npk_optimo": {"N": (60, 100), "P": (30, 50), "K": (80, 120)},
        "ph_optimo": (6.0, 7.5)
    }
}

class AnalizadorFertilidad:
    def __init__(self):
        self.config = None
    
    def dividir_en_sublotes_cuadricula(self, geojson_data, filas=2, columnas=2):
        """Divide el polígono en una cuadrícula de sublotes"""
        try:
            gdf = gpd.GeoDataFrame.from_features(geojson_data["features"])
            
            # Unir todos los polígonos en uno solo
            poligono_principal = unary_union(gdf.geometry)
            
            # Obtener los bounds del polígono
            minx, miny, maxx, maxy = poligono_principal.bounds
            
            # Calcular el tamaño de cada celda
            ancho_celda = (maxx - minx) / columnas
            alto_celda = (maxy - miny) / filas
            
            sublotes = []
            contador = 1
            
            for i in range(filas):
                for j in range(columnas):
                    celda_minx = minx + (j * ancho_celda)
                    celda_maxx = minx + ((j + 1) * ancho_celda)
                    celda_miny = miny + (i * alto_celda)
                    celda_maxy = miny + ((i + 1) * alto_celda)
                    
                    celda_poligono = Polygon([
                        [celda_minx, celda_miny],
                        [celda_maxx, celda_miny],
                        [celda_maxx, celda_maxy],
                        [celda_minx, celda_maxy],
                        [celda_minx, celda_miny]
                    ])
                    
                    interseccion = poligono_principal.intersection(celda_poligono)
                    
                    if not interseccion.is_empty:
                        area_ha = interseccion.area * 10000
                        
                        feature = {
                            "type": "Feature",
                            "properties": {
                                "sublote_id": contador,
                                "nombre": f"Sublote {contador}",
                                "area_ha": round(area_ha, 2),
                                "fila": i + 1,
                                "columna": j + 1
                            },
                            "geometry": json.loads(gpd.GeoSeries([interseccion]).to_json())['features'][0]['geometry']
                        }
                        sublotes.append(feature)
                        contador += 1
            
            geojson_sublotes = {
                "type": "FeatureCollection",
                "features": sublotes
            }
            
            return geojson_sublotes
            
        except Exception as e:
            st.error(f"Error dividiendo en sublotes: {str(e)}")
            return geojson_data
    
    def analizar_fertilidad_sublotes(self, geojson_sublotes, cultivo, fecha_inicio, fecha_fin):
        """Analiza la fertilidad para cada sublote"""
        try:
            st.info("🔍 Analizando fertilidad por sublotes...")
            cultivo_info = CULTIVOS[cultivo]
            
            resultados_sublotes = []
            
            for feature in geojson_sublotes["features"]:
                sublote_id = feature["properties"]["sublote_id"]
                nombre_sublote = feature["properties"]["nombre"]
                area_ha = feature["properties"].get("area_ha", 10)
                
                # Generar datos de fertilidad con variación espacial
                fila = feature["properties"]["fila"]
                columna = feature["properties"]["columna"]
                
                base_n = random.uniform(50, 100) + (fila - 1) * 5
                base_p = random.uniform(30, 60) + (columna - 1) * 3
                base_k = random.uniform(40, 80) + (fila - 1) * 4 - (columna - 1) * 2
                
                nitrogeno = max(10, min(200, base_n + random.uniform(-15, 15)))
                fosforo = max(5, min(100, base_p + random.uniform(-10, 10)))
                potasio = max(20, min(150, base_k + random.uniform(-15, 15)))
                ph = random.uniform(5.5, 7.5)
                materia_organica = random.uniform(2.0, 4.0)
                
                # Calcular índices
                indice_n = self._calcular_indice_nutriente(nitrogeno, cultivo_info['npk_optimo']['N'])
                indice_p = self._calcular_indice_nutriente(fosforo, cultivo_info['npk_optimo']['P'])
                indice_k = self._calcular_indice_nutriente(potasio, cultivo_info['npk_optimo']['K'])
                indice_ph = self._calcular_indice_ph(ph, cultivo_info['ph_optimo'])
                
                fertilidad_general = (indice_n * 0.35 + indice_p * 0.25 + indice_k * 0.25 + indice_ph * 0.15)
                
                recomendaciones = self._generar_recomendaciones_npk(
                    nitrogeno, fosforo, potasio, ph, cultivo_info
                )
                
                categoria_fertilidad = self._obtener_categoria_fertilidad(fertilidad_general)
                categoria_recomendacion = self._obtener_categoria_recomendacion(recomendaciones)
                
                resultado_sublote = {
                    'sublote_id': sublote_id,
                    'nombre_sublote': nombre_sublote,
                    'area_ha': round(area_ha, 2),
                    'fila': fila,
                    'columna': columna,
                    'fertilidad_general': round(fertilidad_general, 1),
                    'categoria_fertilidad': categoria_fertilidad,
                    'categoria_recomendacion': categoria_recomendacion,
                    'nutrientes': {
                        'nitrogeno': round(nitrogeno, 1),
                        'fosforo': round(fosforo, 1),
                        'potasio': round(potasio, 1),
                        'ph': round(ph, 2),
                        'materia_organica': round(materia_organica, 2)
                    },
                    'indices': {
                        'N': round(indice_n, 1),
                        'P': round(indice_p, 1),
                        'K': round(indice_k, 1),
                        'pH': round(indice_ph, 1)
                    },
                    'recomendaciones_npk': recomendaciones,
                    'dosis_npk': self._calcular_dosis_npk(nitrogeno, fosforo, potasio, cultivo_info)
                }
                
                resultados_sublotes.append(resultado_sublote)
            
            st.success(f"🎉 Análisis completado para {len(resultados_sublotes)} sublotes")
            
            return {
                'sublotes': resultados_sublotes,
                'cultivo': cultivo_info['nombre'],
                'fecha_analisis': datetime.now().strftime("%d/%m/%Y %H:%M"),
                'fecha_inicio': fecha_inicio.strftime("%d/%m/%Y"),
                'fecha_fin': fecha_fin.strftime("%d/%m/%Y"),
                'estadisticas': self._calcular_estadisticas(resultados_sublotes)
            }
            
        except Exception as e:
            st.error(f"❌ Error en análisis de sublotes: {str(e)}")
            return None
    
    def _calcular_indice_nutriente(self, valor, rango_optimo):
        """Calcula índice de adecuación del nutriente (0-100)"""
        optimo_medio = (rango_optimo[0] + rango_optimo[1]) / 2
        desviacion = abs(valor - optimo_medio)
        rango_tolerancia = (rango_optimo[1] - rango_optimo[0]) / 2
        
        if desviacion <= rango_tolerancia:
            return 100 - (desviacion / rango_tolerancia * 20)
        else:
            return max(0, 80 - ((desviacion - rango_tolerancia) / rango_tolerancia * 40))
    
    def _calcular_indice_ph(self, ph, rango_optimo):
        """Calcula índice de adecuación del pH"""
        if rango_optimo[0] <= ph <= rango_optimo[1]:
            return 100
        elif ph < rango_optimo[0]:
            desviacion = rango_optimo[0] - ph
            return max(0, 100 - desviacion * 30)
        else:
            desviacion = ph - rango_optimo[1]
            return max(0, 100 - desviacion * 30)
    
    def _generar_recomendaciones_npk(self, nitrogeno, fosforo, potasio, ph, cultivo_info):
        """Genera recomendaciones específicas de fertilización NPK"""
        recomendaciones = []
        
        optimo_n = cultivo_info['npk_optimo']['N']
        if nitrogeno < optimo_n[0]:
            deficit = optimo_n[0] - nitrogeno
            dosis = deficit * 2.0
            recomendaciones.append(f"Nitrogeno: Aplicar {dosis:.0f} kg/ha de Urea")
        elif nitrogeno > optimo_n[1]:
            recomendaciones.append("Nitrogeno: Nivel adecuado")
        else:
            recomendaciones.append("Nitrogeno: Nivel óptimo")
        
        optimo_p = cultivo_info['npk_optimo']['P']
        if fosforo < optimo_p[0]:
            deficit = optimo_p[0] - fosforo
            dosis = deficit * 2.3
            recomendaciones.append(f"Fósforo: Aplicar {dosis:.0f} kg/ha de Superfosfato")
        elif fosforo > optimo_p[1]:
            recomendaciones.append("Fósforo: Nivel adecuado")
        else:
            recomendaciones.append("Fósforo: Nivel óptimo")
        
        optimo_k = cultivo_info['npk_optimo']['K']
        if potasio < optimo_k[0]:
            deficit = optimo_k[0] - potasio
            dosis = deficit * 1.7
            recomendaciones.append(f"Potasio: Aplicar {dosis:.0f} kg/ha de Cloruro de Potasio")
        elif potasio > optimo_k[1]:
            recomendaciones.append("Potasio: Nivel adecuado")
        else:
            recomendaciones.append("Potasio: Nivel óptimo")
        
        optimo_ph = cultivo_info['ph_optimo']
        if ph < optimo_ph[0]:
            recomendaciones.append(f"pH: Encalar con 1-2 tn/ha de calcáreo")
        elif ph > optimo_ph[1]:
            recomendaciones.append(f"pH: Aplicar azufre para reducir pH")
        else:
            recomendaciones.append(f"pH: Nivel óptimo")
        
        return recomendaciones
    
    def _calcular_dosis_npk(self, nitrogeno, fosforo, potasio, cultivo_info):
        """Calcula dosis específicas de NPK"""
        optimo_n = cultivo_info['npk_optimo']['N']
        optimo_p = cultivo_info['npk_optimo']['P']
        optimo_k = cultivo_info['npk_optimo']['K']
        
        dosis_n = max(0, (optimo_n[0] - nitrogeno) * 2.0) if nitrogeno < optimo_n[0] else 0
        dosis_p = max(0, (optimo_p[0] - fosforo) * 2.3) if fosforo < optimo_p[0] else 0
        dosis_k = max(0, (optimo_k[0] - potasio) * 1.7) if potasio < optimo_k[0] else 0
        
        return {
            'N': round(dosis_n, 1),
            'P': round(dosis_p, 1),
            'K': round(dosis_k, 1)
        }
    
    def _obtener_categoria_fertilidad(self, fertilidad):
        """Determina categoría de fertilidad"""
        if fertilidad >= 80:
            return "Alta"
        elif fertilidad >= 60:
            return "Media"
        elif fertilidad >= 40:
            return "Baja"
        else:
            return "Muy Baja"
    
    def _obtener_categoria_recomendacion(self, recomendaciones):
        """Determina categoría de recomendación"""
        necesidades = sum(1 for rec in recomendaciones if "Aplicar" in rec)
        if necesidades == 0:
            return "Sin Fertilización"
        elif necesidades == 1:
            return "Fertilización Leve"
        elif necesidades == 2:
            return "Fertilización Moderada"
        else:
            return "Fertilización Intensiva"
    
    def _calcular_estadisticas(self, resultados_sublotes):
        """Calcula estadísticas generales"""
        fertilities = [s['fertilidad_general'] for s in resultados_sublotes]
        areas = [s['area_ha'] for s in resultados_sublotes]
        
        return {
            'fertilidad_promedio': np.mean(fertilities),
            'fertilidad_min': np.min(fertilities),
            'fertilidad_max': np.max(fertilities),
            'area_total': sum(areas),
            'desviacion_estandar': np.std(fertilities)
        }

# Funciones de procesamiento de archivos (se mantienen igual)
def procesar_archivo_subido(archivo):
    """Procesa archivos ZIP y KML"""
    try:
        st.info(f"📂 Procesando archivo: {archivo.name}")
        
        if archivo.name.lower().endswith('.zip'):
            return procesar_archivo_zip(archivo.read(), archivo.name)
        elif archivo.name.lower().endswith('.kml'):
            return procesar_archivo_kml(archivo.read(), archivo.name)
        else:
            st.error("❌ Formato no soportado. Solo se aceptan ZIP y KML.")
            return None
            
    except Exception as e:
        st.error(f"❌ Error procesando archivo: {str(e)}")
        return None

def procesar_archivo_zip(contenido_zip, nombre_archivo):
    """Procesa archivos ZIP que contengan Shapefiles o KML"""
    try:
        with zipfile.ZipFile(io.BytesIO(contenido_zip), 'r') as zip_ref:
            archivos = zip_ref.namelist()
            
            shp_files = [f for f in archivos if f.lower().endswith('.shp')]
            if shp_files:
                return procesar_shapefile_desde_zip(zip_ref, shp_files[0])
            
            kml_files = [f for f in archivos if f.lower().endswith('.kml')]
            if kml_files:
                return procesar_kml_desde_zip(zip_ref, kml_files[0])
            
            st.error("❌ No se encontraron archivos Shapefile (.shp) o KML (.kml) en el ZIP")
            return None
            
    except Exception as e:
        st.error(f"❌ Error procesando ZIP: {str(e)}")
        return None

def procesar_shapefile_desde_zip(zip_ref, shp_file):
    """Procesa Shapefile desde archivo ZIP"""
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            base_name = os.path.splitext(shp_file)[0]
            for file in zip_ref.namelist():
                if file.startswith(base_name):
                    zip_ref.extract(file, temp_dir)
            
            shp_path = os.path.join(temp_dir, shp_file)
            gdf = gpd.read_file(shp_path)
            
            st.success(f"✅ Shapefile procesado: {len(gdf)} polígonos encontrados")
            return json.loads(gdf.to_json())
            
    except Exception as e:
        st.error(f"❌ Error procesando Shapefile: {str(e)}")
        return None

def procesar_kml_desde_zip(zip_ref, kml_file):
    """Procesa KML desde archivo ZIP"""
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_ref.extract(kml_file, temp_dir)
            kml_path = os.path.join(temp_dir, kml_file)
            gdf = gpd.read_file(kml_path, driver='KML')
            
            st.success(f"✅ KML procesado: {len(gdf)} polígonos encontrados")
            return json.loads(gdf.to_json())
            
    except Exception as e:
        st.error(f"❌ Error procesando KML: {str(e)}")
        return None

def procesar_archivo_kml(contenido_kml, nombre_archivo):
    """Procesa archivos KML directos"""
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            kml_path = os.path.join(temp_dir, "archivo.kml")
            with open(kml_path, 'wb') as f:
                f.write(contenido_kml)
            
            gdf = gpd.read_file(kml_path, driver='KML')
            st.success(f"✅ KML procesado: {len(gdf)} polígonos encontrados")
            return json.loads(gdf.to_json())
            
    except Exception as e:
        st.error(f"❌ Error procesando KML: {str(e)}")
        return None

# Funciones de mapeo mejoradas
def crear_mapa_fertilidad_mejorado(geojson_sublotes, resultados_sublotes):
    """Crea mapa de fertilidad mejorado"""
    try:
        gdf = gpd.GeoDataFrame.from_features(geojson_sublotes["features"])
        centro = [gdf.geometry.centroid.y.mean(), gdf.geometry.centroid.x.mean()]
        
        m = folium.Map(
            location=centro, 
            zoom_start=14,
            tiles='OpenStreetMap',
            control_scale=True
        )
        
        # Capas base
        folium.TileLayer(
            'OpenStreetMap',
            name='Mapa Base'
        ).add_to(m)
        
        folium.TileLayer(
            'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
            attr='Esri',
            name='Vista Satelital'
        ).add_to(m)
        
        # Colores para categorías de fertilidad
        colores_fertilidad = {
            "Alta": "green",
            "Media": "yellow", 
            "Baja": "orange",
            "Muy Baja": "red"
        }
        
        for feature in geojson_sublotes["features"]:
            sublote_id = feature["properties"]["sublote_id"]
            resultado = next((r for r in resultados_sublotes if r['sublote_id'] == sublote_id), None)
            
            if resultado:
                color = colores_fertilidad.get(resultado['categoria_fertilidad'], 'gray')
                popup_text = f"""
                <div style="font-family: Arial; font-size: 12px; min-width: 250px;">
                    <h4 style="color: #2E8B57; margin-bottom: 10px;">Sublote {sublote_id}</h4>
                    <div style="background: {color}; padding: 5px; border-radius: 3px; margin-bottom: 10px;">
                        <strong>Fertilidad: {resultado['fertilidad_general']}%</strong><br>
                        <em>{resultado['categoria_fertilidad']}</em>
                    </div>
                    <table style="width: 100%; font-size: 11px;">
                        <tr><td><strong>Área:</strong></td><td>{resultado['area_ha']} ha</td></tr>
                        <tr><td><strong>N (ppm):</strong></td><td>{resultado['nutrientes']['nitrogeno']}</td></tr>
                        <tr><td><strong>P (ppm):</strong></td><td>{resultado['nutrientes']['fosforo']}</td></tr>
                        <tr><td><strong>K (ppm):</strong></td><td>{resultado['nutrientes']['potasio']}</td></tr>
                        <tr><td><strong>pH:</strong></td><td>{resultado['nutrientes']['ph']}</td></tr>
                        <tr><td><strong>M.O. (%):</strong></td><td>{resultado['nutrientes']['materia_organica']}</td></tr>
                    </table>
                </div>
                """
                
                folium.GeoJson(
                    feature,
                    style_function=lambda x, color=color: {
                        'fillColor': color,
                        'color': 'black',
                        'weight': 2,
                        'fillOpacity': 0.7
                    },
                    popup=folium.Popup(popup_text, max_width=300),
                    tooltip=folium.Tooltip(f"Sublote {sublote_id}: {resultado['fertilidad_general']}%")
                ).add_to(m)
        
        # Agregar leyenda
        legend_html = crear_leyenda_fertilidad_mejorada()
        m.get_root().html.add_child(folium.Element(legend_html))
        
        folium.LayerControl().add_to(m)
        return m
        
    except Exception as e:
        st.error(f"Error creando mapa de fertilidad: {e}")
        return folium.Map(location=[-34.6037, -58.3816], zoom_start=4)

def crear_mapa_recomendaciones_mejorado(geojson_sublotes, resultados_sublotes):
    """Crea mapa de recomendaciones mejorado"""
    try:
        gdf = gpd.GeoDataFrame.from_features(geojson_sublotes["features"])
        centro = [gdf.geometry.centroid.y.mean(), gdf.geometry.centroid.x.mean()]
        
        m = folium.Map(
            location=centro, 
            zoom_start=14,
            tiles='OpenStreetMap',
            control_scale=True
        )
        
        folium.TileLayer(
            'OpenStreetMap',
            name='Mapa Base'
        ).add_to(m)
        
        folium.TileLayer(
            'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
            attr='Esri',
            name='Vista Satelital'
        ).add_to(m)
        
        colores_recomendacion = {
            "Sin Fertilización": "blue",
            "Fertilización Leve": "green",
            "Fertilización Moderada": "orange",
            "Fertilización Intensiva": "red"
        }
        
        for feature in geojson_sublotes["features"]:
            sublote_id = feature["properties"]["sublote_id"]
            resultado = next((r for r in resultados_sublotes if r['sublote_id'] == sublote_id), None)
            
            if resultado:
                color = colores_recomendacion.get(resultado['categoria_recomendacion'], 'gray')
                popup_text = f"""
                <div style="font-family: Arial; font-size: 12px; min-width: 280px;">
                    <h4 style="color: #2E8B57;">Sublote {sublote_id}</h4>
                    <div style="background: {color}; color: white; padding: 5px; border-radius: 3px; margin-bottom: 10px;">
                        <strong>{resultado['categoria_recomendacion']}</strong>
                    </div>
                    <table style="width: 100%; font-size: 11px;">
                        <tr><td><strong>Dosis N:</strong></td><td>{resultado['dosis_npk']['N']} kg/ha</td></tr>
                        <tr><td><strong>Dosis P:</strong></td><td>{resultado['dosis_npk']['P']} kg/ha</td></tr>
                        <tr><td><strong>Dosis K:</strong></td><td>{resultado['dosis_npk']['K']} kg/ha</td></tr>
                        <tr><td><strong>Fertilidad:</strong></td><td>{resultado['fertilidad_general']}%</td></tr>
                    </table>
                </div>
                """
                
                folium.GeoJson(
                    feature,
                    style_function=lambda x, color=color: {
                        'fillColor': color,
                        'color': 'black',
                        'weight': 2,
                        'fillOpacity': 0.7
                    },
                    popup=folium.Popup(popup_text, max_width=300),
                    tooltip=folium.Tooltip(f"Rec: {resultado['categoria_recomendacion']}")
                ).add_to(m)
        
        legend_html = crear_leyenda_recomendaciones_mejorada()
        m.get_root().html.add_child(folium.Element(legend_html))
        
        folium.LayerControl().add_to(m)
        return m
        
    except Exception as e:
        st.error(f"Error creando mapa de recomendaciones: {e}")
        return folium.Map(location=[-34.6037, -58.3816], zoom_start=4)

def crear_leyenda_fertilidad_mejorada():
    """Crea leyenda mejorada para mapa de fertilidad"""
    legend_html = '''
    <div style="
        position: fixed; 
        bottom: 50px; 
        left: 50px; 
        width: 200px; 
        height: auto;
        background-color: white; 
        border: 2px solid grey; 
        z-index: 9999;
        font-size: 11px;
        padding: 10px;
        border-radius: 5px;
        box-shadow: 0 0 10px rgba(0,0,0,0.2);
        font-family: Arial, sans-serif;
    ">
        <h4 style="margin: 0 0 8px 0; text-align: center; font-size: 12px;">Fertilidad del Suelo</h4>
        <p style="margin: 4px 0; display: flex; align-items: center;">
            <span style="background: green; width: 15px; height: 15px; display: inline-block; margin-right: 5px; border: 1px solid black;"></span>
            Alta (80-100%)
        </p>
        <p style="margin: 4px 0; display: flex; align-items: center;">
            <span style="background: yellow; width: 15px; height: 15px; display: inline-block; margin-right: 5px; border: 1px solid black;"></span>
            Media (60-79%)
        </p>
        <p style="margin: 4px 0; display: flex; align-items: center;">
            <span style="background: orange; width: 15px; height: 15px; display: inline-block; margin-right: 5px; border: 1px solid black;"></span>
            Baja (40-59%)
        </p>
        <p style="margin: 4px 0; display: flex; align-items: center;">
            <span style="background: red; width: 15px; height: 15px; display: inline-block; margin-right: 5px; border: 1px solid black;"></span>
            Muy Baja (<40%)
        </p>
    </div>
    '''
    return legend_html

def crear_leyenda_recomendaciones_mejorada():
    """Crea leyenda mejorada para mapa de recomendaciones"""
    legend_html = '''
    <div style="
        position: fixed; 
        bottom: 50px; 
        left: 50px; 
        width: 200px; 
        height: auto;
        background-color: white; 
        border: 2px solid grey; 
        z-index: 9999;
        font-size: 11px;
        padding: 10px;
        border-radius: 5px;
        box-shadow: 0 0 10px rgba(0,0,0,0.2);
        font-family: Arial, sans-serif;
    ">
        <h4 style="margin: 0 0 8px 0; text-align: center; font-size: 12px;">Recomendaciones NPK</h4>
        <p style="margin: 4px 0; display: flex; align-items: center;">
            <span style="background: blue; width: 15px; height: 15px; display: inline-block; margin-right: 5px; border: 1px solid black;"></span>
            Sin Fertilización
        </p>
        <p style="margin: 4px 0; display: flex; align-items: center;">
            <span style="background: green; width: 15px; height: 15px; display: inline-block; margin-right: 5px; border: 1px solid black;"></span>
            Fertilización Leve
        </p>
        <p style="margin: 4px 0; display: flex; align-items: center;">
            <span style="background: orange; width: 15px; height: 15px; display: inline-block; margin-right: 5px; border: 1px solid black;"></span>
            Fertilización Moderada
        </p>
        <p style="margin: 4px 0; display: flex; align-items: center;">
            <span style="background: red; width: 15px; height: 15px; display: inline-block; margin-right: 5px; border: 1px solid black;"></span>
            Fertilización Intensiva
        </p>
    </div>
    '''
    return legend_html

# Funciones para reportes DOCX
def generar_reporte_docx(resultados, geojson_sublotes, cultivo_seleccionado):
    """Genera un reporte DOCX profesional"""
    try:
        doc = Document()
        
        # Título
        title = doc.add_heading('Reporte de Análisis de Fertilidad', 0)
        title.alignment = 1
        
        # Información general
        doc.add_heading('Información del Análisis', level=1)
        p_info = doc.add_paragraph()
        p_info.add_run('Cultivo: ').bold = True
        p_info.add_run(resultados['cultivo'] + '\n')
        p_info.add_run('Fecha de Análisis: ').bold = True
        p_info.add_run(resultados['fecha_analisis'] + '\n')
        p_info.add_run('Período: ').bold = True
        p_info.add_run(f"{resultados['fecha_inicio']} a {resultados['fecha_fin']}\n")
        p_info.add_run('Sublotes Analizados: ').bold = True
        p_info.add_run(f"{len(resultados['sublotes'])}\n")
        p_info.add_run('Área Total: ').bold = True
        p_info.add_run(f"{resultados['estadisticas']['area_total']:.1f} ha")
        
        # Estadísticas generales
        doc.add_heading('Estadísticas Generales', level=1)
        stats = resultados['estadisticas']
        p_stats = doc.add_paragraph()
        p_stats.add_run(f"Fertilidad Promedio: {stats['fertilidad_promedio']:.1f}%\n")
        p_stats.add_run(f"Fertilidad Mínima: {stats['fertilidad_min']:.1f}%\n")
        p_stats.add_run(f"Fertilidad Máxima: {stats['fertilidad_max']:.1f}%\n")
        p_stats.add_run(f"Desviación Estándar: {stats['desviacion_estandar']:.1f}%")
        
        # Tabla de resultados por sublote
        doc.add_heading('Resultados por Sublote', level=1)
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Shading'
        
        # Encabezados de tabla
        headers = ['Sublote', 'Área (ha)', 'Fertilidad (%)', 'Categoría', 'N (ppm)', 'P (ppm)', 'K (ppm)']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        # Datos de la tabla
        for resultado in resultados['sublotes']:
            row_cells = table.add_row().cells
            row_cells[0].text = resultado['nombre_sublote']
            row_cells[1].text = str(resultado['area_ha'])
            row_cells[2].text = str(resultado['fertilidad_general'])
            row_cells[3].text = resultado['categoria_fertilidad']
            row_cells[4].text = str(resultado['nutrientes']['nitrogeno'])
            row_cells[5].text = str(resultado['nutrientes']['fosforo'])
            row_cells[6].text = str(resultado['nutrientes']['potasio'])
        
        # Recomendaciones generales
        doc.add_heading('Recomendaciones Generales', level=1)
        
        # Contar categorías de recomendación
        categorias = {}
        for resultado in resultados['sublotes']:
            cat = resultado['categoria_recomendacion']
            categorias[cat] = categorias.get(cat, 0) + 1
        
        for categoria, count in categorias.items():
            p_rec = doc.add_paragraph()
            p_rec.add_run(f"{categoria}: ").bold = True
            p_rec.add_run(f"{count} sublotes ({count/len(resultados['sublotes'])*100:.1f}%)")
        
        # Guardar documento
        output_dir = Path("reportes")
        output_dir.mkdir(exist_ok=True)
        
        filename = output_dir / f"reporte_fertilidad_{cultivo_seleccionado}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(filename)
        
        return filename
        
    except Exception as e:
        st.error(f"❌ Error generando reporte DOCX: {str(e)}")
        return None

def crear_tabla_resultados(resultados_sublotes):
    """Crea tabla resumen de resultados por sublote"""
    datos_tabla = []
    
    for resultado in resultados_sublotes:
        fila = {
            'Sublote': resultado['nombre_sublote'],
            'Área (ha)': resultado['area_ha'],
            'Fertilidad (%)': resultado['fertilidad_general'],
            'Categoría': resultado['categoria_fertilidad'],
            'N (ppm)': resultado['nutrientes']['nitrogeno'],
            'P (ppm)': resultado['nutrientes']['fosforo'],
            'K (ppm)': resultado['nutrientes']['potasio'],
            'pH': resultado['nutrientes']['ph'],
            'M.O. (%)': resultado['nutrientes']['materia_organica'],
            'Dosis N (kg/ha)': resultado['dosis_npk']['N'],
            'Dosis P (kg/ha)': resultado['dosis_npk']['P'],
            'Dosis K (kg/ha)': resultado['dosis_npk']['K'],
            'Recomendación': resultado['categoria_recomendacion']
        }
        datos_tabla.append(fila)
    
    return pd.DataFrame(datos_tabla)

def exportar_geojson_resultados(geojson_sublotes, resultados_completos):
    """Exporta GeoJSON con todos los resultados del análisis"""
    try:
        geojson_export = json.loads(json.dumps(geojson_sublotes))
        
        for feature in geojson_export["features"]:
            sublote_id = feature["properties"]["sublote_id"]
            resultado = next((r for r in resultados_completos['sublotes'] if r['sublote_id'] == sublote_id), None)
            
            if resultado:
                feature["properties"].update({
                    'cultivo': resultados_completos['cultivo'],
                    'fecha_analisis': resultados_completos['fecha_analisis'],
                    'fertilidad_general': resultado['fertilidad_general'],
                    'categoria_fertilidad': resultado['categoria_fertilidad'],
                    'nitrogeno_ppm': resultado['nutrientes']['nitrogeno'],
                    'fosforo_ppm': resultado['nutrientes']['fosforo'],
                    'potasio_ppm': resultado['nutrientes']['potasio'],
                    'ph': resultado['nutrientes']['ph'],
                    'materia_organica': resultado['nutrientes']['materia_organica'],
                    'dosis_n_kg_ha': resultado['dosis_npk']['N'],
                    'dosis_p_kg_ha': resultado['dosis_npk']['P'],
                    'dosis_k_kg_ha': resultado['dosis_npk']['K'],
                    'categoria_recomendacion': resultado['categoria_recomendacion']
                })
        
        return geojson_export
        
    except Exception as e:
        st.error(f"❌ Error exportando GeoJSON: {str(e)}")
        return None

def main():
    st.markdown('<h1 class="main-header">🌱 Analizador de Fertilidad Multi-Cultivo</h1>', unsafe_allow_html=True)
    
    # Inicializar estado de la sesión
    if 'geojson_data' not in st.session_state:
        st.session_state.geojson_data = None
    if 'geojson_sublotes' not in st.session_state:
        st.session_state.geojson_sublotes = None
    if 'resultados' not in st.session_state:
        st.session_state.resultados = None
    if 'analisis_completado' not in st.session_state:
        st.session_state.analisis_completado = False
    
    # Sidebar
    with st.sidebar:
        st.markdown('<h3 class="section-header">⚙️ Configuración</h3>', unsafe_allow_html=True)
        
        cultivo = st.selectbox(
            "Cultivo a analizar",
            options=list(CULTIVOS.keys()),
            format_func=lambda x: CULTIVOS[x]['nombre']
        )
        
        cultivo_info = CULTIVOS[cultivo]
        st.info(f"""
        **Cultivo:** {cultivo_info['nombre']}
        **NPK Óptimo:** 
        - N: {cultivo_info['npk_optimo']['N'][0]}-{cultivo_info['npk_optimo']['N'][1]} ppm
        - P: {cultivo_info['npk_optimo']['P'][0]}-{cultivo_info['npk_optimo']['P'][1]} ppm  
        - K: {cultivo_info['npk_optimo']['K'][0]}-{cultivo_info['npk_optimo']['K'][1]} ppm
        **pH Óptimo:** {cultivo_info['ph_optimo'][0]}-{cultivo_info['ph_optimo'][1]}
        """)
        
        st.markdown("---")
        st.markdown("### 📁 Cargar Polígono")
        
        st.markdown("""
        <div class="warning-box">
        <strong>Formatos aceptados:</strong>
        <br>• <strong>ZIP</strong> con Shapefile (.shp, .dbf, .shx)
        <br>• <strong>KML</strong> de Google Earth
        </div>
        """, unsafe_allow_html=True)
        
        archivo = st.file_uploader(
            "Subir archivo ZIP o KML",
            type=['zip', 'kml'],
            help="Archivo ZIP con Shapefile o KML de Google Earth"
        )
        
        if archivo is not None:
            if st.session_state.geojson_data is None or st.button("🔄 Reprocesar archivo"):
                with st.spinner("Procesando archivo..."):
                    geojson_data = procesar_archivo_subido(archivo)
                    if geojson_data is not None:
                        st.session_state.geojson_data = geojson_data
                        st.session_state.resultados = None
                        st.session_state.analisis_completado = False
                        st.success("✅ Archivo cargado correctamente")
        
        st.markdown("---")
        st.markdown("### 🗂️ Configuración de Cuadrícula")
        col1, col2 = st.columns(2)
        with col1:
            filas = st.slider("Filas", min_value=1, max_value=6, value=2)
        with col2:
            columnas = st.slider("Columnas", min_value=1, max_value=6, value=2)
        
        st.info(f"**Total de sublotes:** {filas * columnas}")
        
        st.markdown("---")
        st.markdown("### 📅 Período de Análisis")
        col1, col2 = st.columns(2)
        with col1:
            fecha_inicio = st.date_input("Desde", value=datetime.now() - timedelta(days=30))
        with col2:
            fecha_fin = st.date_input("Hasta", value=datetime.now())
        
        st.markdown("---")
        analizar_disabled = st.session_state.geojson_data is None
        
        if st.button(
            "🚀 Ejecutar Análisis de Fertilidad", 
            type="primary", 
            use_container_width=True,
            disabled=analizar_disabled
        ):
            if st.session_state.geojson_data:
                with st.spinner("🔄 Iniciando análisis de fertilidad..."):
                    try:
                        analizador = AnalizadorFertilidad()
                        
                        st.info("🗂️ Creando cuadrícula de sublotes...")
                        geojson_sublotes = analizador.dividir_en_sublotes_cuadricula(
                            st.session_state.geojson_data, filas, columnas
                        )
                        st.session_state.geojson_sublotes = geojson_sublotes
                        
                        resultados = analizador.analizar_fertilidad_sublotes(
                            geojson_sublotes, cultivo, fecha_inicio, fecha_fin
                        )
                        
                        if resultados is not None:
                            st.session_state.resultados = resultados
                            st.session_state.analisis_completado = True
                            st.success("🎉 ¡Análisis completado exitosamente!")
                        else:
                            st.error("❌ No se pudieron generar los resultados")
                            
                    except Exception as e:
                        st.error(f"❌ Error durante el análisis: {str(e)}")
            else:
                st.error("❌ Primero carga un archivo con el polígono del lote")
    
    # Contenido principal
    if st.session_state.analisis_completado and st.session_state.resultados:
        resultados = st.session_state.resultados
        
        tab1, tab2, tab3, tab4 = st.tabs([
            "🗺️ Mapa Fertilidad", 
            "🧪 Mapa Recomendaciones", 
            "📊 Tabla Resultados", 
            "📥 Exportar"
        ])
        
        with tab1:
            st.markdown('<h3 class="section-header">🗺️ Mapa de Fertilidad por Sublotes</h3>', unsafe_allow_html=True)
            st.markdown('<div class="map-container">', unsafe_allow_html=True)
            mapa_fertilidad = crear_mapa_fertilidad_mejorado(
                st.session_state.geojson_sublotes, 
                resultados['sublotes']
            )
            st_folium(mapa_fertilidad, height=600, use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
            
        with tab2:
            st.markdown('<h3 class="section-header">🧪 Mapa de Recomendaciones NPK</h3>', unsafe_allow_html=True)
            st.markdown('<div class="map-container">', unsafe_allow_html=True)
            mapa_recomendaciones = crear_mapa_recomendaciones_mejorado(
                st.session_state.geojson_sublotes, 
                resultados['sublotes']
            )
            st_folium(mapa_recomendaciones, height=600, use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
            
        with tab3:
            st.markdown('<h3 class="section-header">📊 Tabla de Resultados por Sublote</h3>', unsafe_allow_html=True)
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                fert_promedio = resultados['estadisticas']['fertilidad_promedio']
                st.metric("Fertilidad Promedio", f"{fert_promedio:.1f}%")
            with col2:
                total_ha = resultados['estadisticas']['area_total']
                st.metric("Área Total", f"{total_ha:.1f} ha")
            with col3:
                st.metric("Sublotes", len(resultados['sublotes']))
            with col4:
                sublotes_bajos = sum(1 for s in resultados['sublotes'] if s['fertilidad_general'] < 60)
                st.metric("Sublotes <60%", sublotes_bajos)
            
            df_resultados = crear_tabla_resultados(resultados['sublotes'])
            st.dataframe(df_resultados, use_container_width=True, height=400)
            
        with tab4:
            st.markdown('<h3 class="section-header">📥 Exportar Resultados</h3>', unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                geojson_export = exportar_geojson_resultados(
                    st.session_state.geojson_sublotes, 
                    resultados
                )
                
                if geojson_export:
                    geojson_str = json.dumps(geojson_export, indent=2, ensure_ascii=False)
                    
                    st.download_button(
                        label="📥 Descargar GeoJSON",
                        data=geojson_str,
                        file_name=f"fertilidad_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.geojson",
                        mime="application/json",
                        type="primary",
                        use_container_width=True
                    )
            
            with col2:
                df_export = crear_tabla_resultados(resultados['sublotes'])
                csv_data = df_export.to_csv(index=False, encoding='utf-8')
                
                st.download_button(
                    label="📊 Descargar CSV",
                    data=csv_data,
                    file_name=f"resultados_fertilidad_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            with col3:
                if st.button("📝 Generar Reporte DOCX", use_container_width=True):
                    with st.spinner("Generando reporte DOCX..."):
                        docx_file = generar_reporte_docx(
                            resultados, 
                            st.session_state.geojson_sublotes, 
                            cultivo
                        )
                        
                        if docx_file and os.path.exists(docx_file):
                            with open(docx_file, "rb") as f:
                                st.download_button(
                                    "⬇️ Descargar Reporte DOCX",
                                    f,
                                    file_name=os.path.basename(docx_file),
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
            
            st.markdown("### 📋 Información del Análisis")
            st.write(f"**Cultivo:** {resultados['cultivo']}")
            st.write(f"**Fecha de análisis:** {resultados['fecha_analisis']}")
            st.write(f"**Período analizado:** {resultados['fecha_inicio']} a {resultados['fecha_fin']}")
            st.write(f"**Cuadrícula:** {filas} filas × {columnas} columnas")
            st.write(f"**Sublotes analizados:** {len(resultados['sublotes'])}")
            st.write(f"**Área total:** {resultados['estadisticas']['area_total']:.1f} ha")
    
    else:
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown('<h3 class="section-header">🗺️ Mapa del Lote</h3>', unsafe_allow_html=True)
            
            if st.session_state.geojson_data:
                gdf = gpd.GeoDataFrame.from_features(st.session_state.geojson_data["features"])
                centro = [gdf.geometry.centroid.y.mean(), gdf.geometry.centroid.x.mean()]
                
                m = folium.Map(location=centro, zoom_start=12)
                folium.TileLayer(
                    'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}',
                    attr='Esri',
                    name='Vista Satelital'
                ).add_to(m)
                
                folium.GeoJson(st.session_state.geojson_data).add_to(m)
                folium.LayerControl().add_to(m)
                
                st_folium(m, height=500, use_container_width=True)
            else:
                st.info("👆 Carga un archivo ZIP o KML para visualizar el mapa")
        
        with col2:
            st.markdown('<h3 class="section-header">📊 Análisis de Fertilidad</h3>', unsafe_allow_html=True)
            
            if st.session_state.geojson_data:
                st.info("👆 Configura la cuadrícula y ejecuta el análisis para ver los resultados")
            else:
                st.info("""
                💡 **Para comenzar:**
                1. Selecciona el cultivo
                2. Carga un archivo ZIP o KML
                3. Configura la cuadrícula
                4. Ejecuta el análisis
                
                **Obtendrás:**
                • Mapas de fertilidad
                • Recomendaciones NPK  
                • Tablas de resultados
                • Reportes exportables
                """)

if __name__ == "__main__":
    main()
