from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
from pathlib import Path

def generate_report(output_path, metadata, maps, fertility_stats, crop,
                    producers_notes=True, technician_notes=True):
    """
    Genera informe DOCX de fertilidad y recomendaciones agroecológicas.
    Devuelve un diccionario con rutas de salida (docx/pdf/pptx si aplica)
    """

    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = f"{output_path}.docx"

    # Crear documento
    doc = Document()

    # Portada
    doc.add_heading("INFORME DE FERTILIDAD Y RECOMENDACIONES", level=1)
    doc.add_paragraph(f"Proyecto: {metadata.get('proyecto', 'Análisis de Fertilidad')}")
    doc.add_paragraph(f"Usuario: {metadata.get('usuario', 'Técnico responsable')}")
    doc.add_paragraph(f"Fecha: {metadata.get('fecha', datetime.now().strftime('%Y-%m-%d'))}")
    doc.add_paragraph(f"Cultivo analizado: {crop.capitalize()}")
    doc.add_paragraph("\n")

    # Mapas
    doc.add_heading("Mapas de Fertilidad y Recomendaciones", level=2)
    for key, path in maps.items():
        if os.path.exists(path):
            doc.add_paragraph(key.capitalize())
            doc.add_picture(path, width=Inches(5.5))
            doc.add_paragraph("\n")

    # Fertilidad promedio
    doc.add_heading("Resumen de Fertilidad Promedio", level=2)
    table = doc.add_table(rows=2, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "N (kg/ha)"
    hdr[1].text = "P (kg/ha)"
    hdr[2].text = "K (kg/ha)"
    row = table.rows[1].cells
    row[0].text = f"{fertility_stats.get('N', 0):.1f}"
    row[1].text = f"{fertility_stats.get('P', 0):.1f}"
    row[2].text = f"{fertility_stats.get('K', 0):.1f}"

    doc.add_paragraph("\n")

    # Recomendaciones técnicas
    if technician_notes:
        doc.add_heading("Recomendaciones Técnicas (Agrónomos)", level=2)
        doc.add_paragraph(
            "• Ajustar dosis de NPK según análisis de suelo y rendimiento objetivo.\n"
            "• Monitorear pH y materia orgánica para mantener la eficiencia del fertilizante.\n"
            "• Implementar rotaciones de cultivo y cobertura permanente.\n"
            "• Aplicar biofertilizantes o compost en áreas con menor fertilidad.\n"
        )

    # Recomendaciones para productores
    if producers_notes:
        doc.add_heading("Recomendaciones Prácticas (Productores)", level=2)
        doc.add_paragraph(
            "• Evitar sobrelabranza y mantener cobertura vegetal.\n"
            "• Incorporar residuos de cosecha y materia orgánica.\n"
            "• Promover la infiltración del agua mediante manejo del pastoreo.\n"
            "• Mantener franjas de biodiversidad y corredores biológicos.\n"
        )

    # Agricultura regenerativa
    doc.add_heading("Enfoque de Agricultura Regenerativa", level=2)
    doc.add_paragraph(
        "La fertilidad del suelo debe entenderse como un proceso vivo. "
        "El manejo regenerativo busca restaurar la estructura, microbiología y resiliencia del ecosistema agrícola. "
        "Se recomienda priorizar la reducción de insumos sintéticos, "
        "uso de cultivos de cobertura, pastoreo planificado y diversificación de especies."
    )

    # Guardar DOCX
    doc.save(docx_path)

    return {"docx": docx_path}
